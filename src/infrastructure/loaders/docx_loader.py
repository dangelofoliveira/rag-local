from docx import Document as DocxDocument
from src.infrastructure.loaders.base_loader import BaseLoader
from src.domain.entities.document import Document
from src.config.settings import settings


class DocxLoader(BaseLoader):

    def load(self, path: str) -> list[Document]:
        docx = DocxDocument(path)
        text = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])

        chunks = self._make_chunks(text, settings.chunk_size, settings.chunk_overlap)

        return [
            Document(
                id=self._generate_id(path, i, chunk),
                content=chunk,
                source=path,
                chunk_index=i,
                metadata={"type": "docx"}
            )
            for i, chunk in enumerate(chunks)
        ]